from flask import Flask, request, jsonify, send_file
from docx import Document
from datetime import datetime
import os
import platform
import subprocess

app = Flask(__name__)

# Helper function to replace placeholders while preserving formatting
def replace_placeholders(doc, placeholders):
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

def edit_nda_template(template_path, output_path, placeholders):
    try:
        doc = Document(template_path)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")

def convert_to_pdf(doc_path, pdf_path):
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    if platform.system() == "Windows":
        try:
            import comtypes.client
            import pythoncom
            pythoncom.CoInitialize()
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
        except Exception as e:
            raise Exception(f"Error using COM on Windows: {e}")
    else:
        try:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), doc_path],
                check=True
            )
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error using LibreOffice: {e}")

@app.route('/generate-nda', methods=['POST'])
def generate_nda():
    data = request.json

    client_name = data.get("client_name")
    company_name = data.get("company_name")
    address = data.get("address")
    designation = data.get("designation")
    date_field = data.get("date_field", datetime.today().strftime("%Y-%m-%d"))

    placeholders = {
        "<<Client Name>>": client_name,
        "<<Company Name>>": company_name,
        "<<Address>>": address,
        "<<Designation>>": designation,
        "<<Date>>": datetime.strptime(date_field, "%Y-%m-%d").strftime("%d-%m-%Y"),
    }

    template_path = "Non Disclosure Agreement.docx"
    formatted_date = datetime.strptime(date_field, "%Y-%m-%d").strftime("%d %b %Y")
    file_name = f"NDA Agreement - {client_name} {formatted_date}.docx"
    pdf_file_name = f"NDA Agreement - {client_name} {formatted_date}.pdf"
    word_output_path = os.path.join(os.getcwd(), file_name)
    pdf_output_path = os.path.join(os.getcwd(), pdf_file_name)

    try:
        updated_path = edit_nda_template(template_path, word_output_path, placeholders)
        convert_to_pdf(updated_path, pdf_output_path)

        return jsonify({
            "status": "success",
            "word_document": file_name,
            "pdf_document": pdf_file_name
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(os.getcwd(), filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({"status": "error", "message": "File not found!"}), 404

@app.route('/', methods=['GET'])
def home():
    return jsonify({
        "message": "NDA Document Generator API is running!",
        "endpoints": {
            "POST /generate-nda": "Generate an NDA document",
            "GET /download/<filename>": "Download a generated document"
        }
    })

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=8080, debug=True)
